import os
import pdfplumber
import docx
import re
import traceback
import datetime
from docx.shared import Pt, Cm
from openai import OpenAI

# Funciones para extraer texto
def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                content = page.extract_text()
                if content:
                    text += content + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error al extraer texto del PDF: {str(e)}")
        return ""

def extract_text_from_word(docx_path):
    text = ""
    try:
        doc = docx.Document(docx_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        return text.strip()
    except Exception as e:
        print(f"Error al extraer texto del documento Word: {str(e)}")
        return ""

# Prompt adaptado para OpenAI GPT-4
def create_prompt(cv_text):
    return f'''
    Eres un especialista en RRHH y debes extraer TODA la información relevante del siguiente CV, sin omitir ninguna experiencia laboral. Es CRUCIAL que captures todo el historial laboral completo.
    
    Necesito que extraigas:
    
    1. NOMBRE COMPLETO DEL CANDIDATO
    
    2. PROFESIÓN DEL CANDIDATO (exactamente como aparece)
    
    3. RESEÑA PROFESIONAL (exactamente 4 puntos):
       a) Experiencia profesional y áreas de especialización
       b) Roles principales en los que se ha desempeñado
       c) Tipos de proyectos en los que ha trabajado
       d) Habilidades blandas destacadas
    
    4. HISTORIAL LABORAL (IMPORTANTE: extrae TODAS las experiencias, sin omitir ninguna):
       - Cargo y empresa (formato exacto: "Cargo, Empresa")
       - Período completo (formato exacto: "Mes Año -- Mes Año" o "Mes Año -- A la fecha")
       - Cliente (si aplica, formato: "Cliente: Nombre")
       - TODAS las responsabilidades/experiencias mencionadas para cada trabajo
       - Tecnologías utilizadas (separadas por comas)
    
    5. STACK TECNOLÓGICO:
       - Lista COMPLETA de TODAS las tecnologías, plataformas, herramientas y software mencionados en cualquier parte del CV
    
    6. FORMACIÓN ACADÉMICA:
       - Formato: "Año, Título, Institución"
    
    7. CURSOS Y CAPACITACIONES:
       - Formato: "Año, Nombre del curso, Institución"
    
    INSTRUCCIONES CRÍTICAS:
    - NO OMITAS ninguna experiencia laboral. El CV completo debe ser procesado.
    - NO INCLUYAS frases como "(Tipo de experiencia y rubros)" o similar
    - Organiza el historial laboral en orden cronológico inverso (más reciente primero)
    - Incluye TODAS las tecnologías mencionadas en el CV
    - Si no encuentras alguna información específica, indica "No especificado"
    - No inventes información, pero los meses deben ser Enero, Febrero, etc.
    
    CV A ANALIZAR:
    
    {cv_text}
    
    Responde usando este formato exacto:
    
    <NOMBRE>
    Nombre completo
    </NOMBRE>
    
    <PROFESION>
    Profesión/título
    </PROFESION>
    
    <RESENA>
    - Primer punto sobre experiencia
    - Segundo punto sobre roles
    - Tercer punto sobre proyectos
    - Cuarto punto sobre habilidades
    </RESENA>
    
    <HISTORIAL>
    <TRABAJO>
    <CARGO>Cargo, Empresa</CARGO>
    <PERIODO>Período completo (Mes Año - Mes Año)</PERIODO>
    <CLIENTE>Cliente (si aplica)</CLIENTE>
    <RESPONSABILIDADES>
    - Responsabilidad 1
    - Responsabilidad 2
    - Responsabilidad 3
    - Responsabilidad 4
    - ... (incluir TODAS las responsabilidades)
    </RESPONSABILIDADES>
    <TECNOLOGIAS>Tecnología 1, Tecnología 2, etc.</TECNOLOGIAS>
    </TRABAJO>
    
    (Repite para CADA posición, sin omitir ninguna)
    </HISTORIAL>
    
    <STACK>
    - Tecnología 1
    - Tecnología 2
    - Tecnología 3
    - ... (incluir TODAS las tecnologías)
    </STACK>
    
    <FORMACION>
    - Formación 1
    - Formación 2
    </FORMACION>
    
    <CURSOS>
    - Curso 1
    - Curso 2
    </CURSOS>
    '''

# Reemplazar process_with_gemini por process_with_openai
def process_with_openai(cv_text, api_key, model="o4-mini"):
    try:
        print(f"Procesando CV con OpenAI {model}...")
        
        # Inicializar cliente de OpenAI
        client = OpenAI(api_key=api_key)
        
        # Modelos a probar en caso de fallo
        models_to_try = []
        
        if model in models_to_try:
            models_to_try.remove(model)
        models_to_try.insert(0, model)
        
        for model_to_try in models_to_try:
            print(f"Intentando con modelo: {model_to_try}")
            
            try:
                # Realizar la llamada a la API de OpenAI
                response = client.chat.completions.create(
                    model=model_to_try,
                    messages=[
                        {"role": "system", "content": "Eres un asistente especializado en RRHH que extrae información detallada de CVs."},
                        {"role": "user", "content": create_prompt(cv_text)}
                    ]
                )
                
                # Verificar si la respuesta es válida
                if response and hasattr(response, 'choices') and len(response.choices) > 0:
                    return response.choices[0].message.content
                
            except Exception as e:
                print(f"Error con modelo {model_to_try}: {str(e)}")
                continue
        
        print("No se pudo procesar el CV con ninguno de los modelos disponibles.")
        return None
    except Exception as e:
        print(f"Error al procesar con OpenAI: {str(e)}")
        traceback.print_exc()
    return None

# Extraer información de la respuesta formateada
def extract_sections(response_text):
    result = {
        "nombre": "",
        "profesion": "",
        "resena": [],
        "historial": [],
        "stack": [],
        "formacion": [],
        "cursos": []
    }
    
    # Extraer nombre
    nombre_match = re.search(r'<NOMBRE>(.*?)</NOMBRE>', response_text, re.DOTALL)
    if nombre_match:
        result["nombre"] = nombre_match.group(1).strip()
    
    # Extraer profesión
    profesion_match = re.search(r'<PROFESION>(.*?)</PROFESION>', response_text, re.DOTALL)
    if profesion_match:
        result["profesion"] = profesion_match.group(1).strip()
    
    # Extraer reseña profesional
    resena_match = re.search(r'<RESENA>(.*?)</RESENA>', response_text, re.DOTALL)
    if resena_match:
        resena_text = resena_match.group(1).strip()
        for line in resena_text.split('\n'):
            clean_line = line.strip()
            if clean_line and clean_line.startswith('-'):
                result["resena"].append(clean_line[1:].strip())
    
    # Extraer historial laboral
    historial_match = re.search(r'<HISTORIAL>(.*?)</HISTORIAL>', response_text, re.DOTALL)
    if historial_match:
        historial_text = historial_match.group(1).strip()
        
        # Buscar cada bloque de trabajo
        trabajo_blocks = re.finditer(r'<TRABAJO>(.*?)</TRABAJO>', historial_text, re.DOTALL)
        
        for trabajo_match in trabajo_blocks:
            trabajo_text = trabajo_match.group(1).strip()
            
            trabajo = {
                "cargo": "",
                "periodo": "",
                "cliente": "",
                "responsabilidades": [],
                "tecnologias": ""
            }
            
            # Extraer campos específicos
            cargo_match = re.search(r'<CARGO>(.*?)</CARGO>', trabajo_text, re.DOTALL)
            if cargo_match:
                trabajo["cargo"] = cargo_match.group(1).strip()
            
            periodo_match = re.search(r'<PERIODO>(.*?)</PERIODO>', trabajo_text, re.DOTALL)
            if periodo_match:
                trabajo["periodo"] = periodo_match.group(1).strip()
            
            cliente_match = re.search(r'<CLIENTE>(.*?)</CLIENTE>', trabajo_text, re.DOTALL)
            if cliente_match:
                cliente = cliente_match.group(1).strip()
                if cliente.lower() not in ["no aplica", "no especificado", "n/a"]:
                    trabajo["cliente"] = cliente
            
            resp_match = re.search(r'<RESPONSABILIDADES>(.*?)</RESPONSABILIDADES>', trabajo_text, re.DOTALL)
            if resp_match:
                resp_text = resp_match.group(1).strip()
                for line in resp_text.split('\n'):
                    clean_line = line.strip()
                    if clean_line and clean_line.startswith('-'):
                        trabajo["responsabilidades"].append(clean_line[1:].strip())
            
            tech_match = re.search(r'<TECNOLOGIAS>(.*?)</TECNOLOGIAS>', trabajo_text, re.DOTALL)
            if tech_match:
                trabajo["tecnologias"] = tech_match.group(1).strip()
            
            result["historial"].append(trabajo)
    
    # Extraer stack tecnológico
    stack_match = re.search(r'<STACK>(.*?)</STACK>', response_text, re.DOTALL)
    if stack_match:
        stack_text = stack_match.group(1).strip()
        for line in stack_text.split('\n'):
            clean_line = line.strip()
            if clean_line and clean_line.startswith('-'):
                result["stack"].append(clean_line[1:].strip())
    
    # Extraer formación académica
    formacion_match = re.search(r'<FORMACION>(.*?)</FORMACION>', response_text, re.DOTALL)
    if formacion_match:
        formacion_text = formacion_match.group(1).strip()
        for line in formacion_text.split('\n'):
            clean_line = line.strip()
            if clean_line and clean_line.startswith('-'):
                result["formacion"].append(clean_line[1:].strip())
    
    # Extraer cursos y capacitaciones
    cursos_match = re.search(r'<CURSOS>(.*?)</CURSOS>', response_text, re.DOTALL)
    if cursos_match:
        cursos_text = cursos_match.group(1).strip()
        for line in cursos_text.split('\n'):
            clean_line = line.strip()
            if clean_line and clean_line.startswith('-'):
                result["cursos"].append(clean_line[1:].strip())
    
    return result

# Función para reemplazar texto en párrafos manteniendo formato
def replace_paragraph_text(paragraph, new_text):
    # Si el párrafo está vacío o el nuevo texto está vacío, no hacer nada
    if not paragraph.runs or not new_text:
        return
    
    # Guardar las propiedades del primer run (fuente, tamaño, negrita, etc.)
    first_run = paragraph.runs[0]
    font_name = first_run.font.name
    font_size = first_run.font.size
    bold = first_run.bold
    italic = first_run.italic
    
    # Limpiar el párrafo
    paragraph.clear()
    
    # Añadir el nuevo texto con las mismas propiedades
    run = paragraph.add_run(new_text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.italic = italic

# Función para rellenar la plantilla exactamente
def fill_template_exact(cv_data, template_path, output_path):
    try:
        # Cargar plantilla
        doc = docx.Document(template_path)
        
        # 1. REEMPLAZAR NOMBRE Y PROFESIÓN
        for para in doc.paragraphs:
            if "Nombre Candidato" in para.text:
                replace_paragraph_text(para, cv_data["nombre"])
            elif para.text.strip() == "PROFESIÓN":
                replace_paragraph_text(para, cv_data["profesion"].upper())
        
        # 2. REEMPLAZAR RESEÑA PROFESIONAL
        # Buscar los párrafos que contienen el texto de reseña profesional
        resena_paras = []
        for para in doc.paragraphs:
            if any(marker in para.text for marker in [
                "(Tipo de experiencia y rubros)",
                "(3 roles relevantes",
                "(Tipo de proyectos",
                "(Soft skills"
            ]):
                resena_paras.append(para)
        
        # Ordenar los párrafos según el orden estándar
        resena_order = [
            "(Tipo de experiencia y rubros)",
            "(3 roles relevantes",
            "(Tipo de proyectos",
            "(Soft skills"
        ]
        resena_paras.sort(key=lambda p: next((i for i, m in enumerate(resena_order) if m in p.text), 999))
        
        # Reemplazar el texto manteniendo solo el bullet point
        for i, para in enumerate(resena_paras):
            if i < len(cv_data["resena"]):
                # Construir el nuevo texto: solo bullet point + nuevo texto
                para.clear()
                para.add_run(cv_data["resena"][i])
        
        # 3. REEMPLAZAR HISTORIAL LABORAL
        # Encontrar los bloques de experiencia laboral en la plantilla
        historial_blocks = []
        cargo_paragraphs = []
        
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == "Cargo, Empresa":
                para.clear()
                cargo_paragraphs.append(i)
        
        # Extraer bloques de experiencia a partir de los cargos
        for start_idx in cargo_paragraphs:
            end_idx = next((i for i in cargo_paragraphs if i > start_idx), len(doc.paragraphs))
            historial_blocks.append((start_idx, end_idx))
        
        # Reemplazar el contenido de cada bloque
        for i, (start_idx, end_idx) in enumerate(historial_blocks):
            if i < len(cv_data["historial"]):
                trabajo = cv_data["historial"][i]
                
                # Cargo, Empresa (en cursiva y negrita, sin palabra "Cargo")
                para = doc.paragraphs[start_idx]
                para.clear()
                run = para.add_run(trabajo["cargo"])
                run.bold = True
                run.italic = True
                run.font.size = Pt(12)
                
                # Fecha
                fecha_idx = start_idx + 1
                if fecha_idx < end_idx:
                    replace_paragraph_text(doc.paragraphs[fecha_idx], trabajo["periodo"])
                
                # Cliente (puede o no existir)
                cliente_idx = fecha_idx + 1
                if cliente_idx < end_idx and "Cliente:" in doc.paragraphs[cliente_idx].text:
                    if trabajo.get("cliente"):
                        # Mantener "Cliente:" en negrita
                        para = doc.paragraphs[cliente_idx]
                        para.clear()
                        cliente_run = para.add_run("Cliente: ")
                        cliente_run.bold = True
                        para.add_run(trabajo["cliente"])
                    else:
                        # Si no hay cliente, eliminamos ese párrafo
                        p = doc.paragraphs[cliente_idx]._element
                        p.getparent().remove(p)
                        p._p = p._element = None
                        # Ajustar los índices
                        cliente_idx -= 1
                        end_idx -= 1
                else:
                    cliente_idx = fecha_idx
                
                # Experiencias (Responsabilidades) - solo bullet cuadrado
                exp_start_idx = cliente_idx + 1
                exp_count = 0
                
                for j in range(4):  # La plantilla tiene 4 puntos de experiencia
                    exp_idx = exp_start_idx + j
                    if exp_idx < end_idx and "Experiencia Laboral" in doc.paragraphs[exp_idx].text:
                        para = doc.paragraphs[exp_idx]
                        para.clear()
                        
                        # Texto de la responsabilidad
                        if j < len(trabajo["responsabilidades"]):
                            para.add_run(trabajo["responsabilidades"][j])
                        else:
                            para.add_run("No especificado")
                        
                        exp_count += 1
                
                # Tecnologías utilizadas
                tech_idx = exp_start_idx + exp_count
                if tech_idx < end_idx and "Tecnologías utilizadas:" in doc.paragraphs[tech_idx].text:
                    para = doc.paragraphs[tech_idx]
                    para.clear()
                    tech_run = para.add_run("Tecnologías utilizadas: ")
                    tech_run.bold = True
                    para.add_run(trabajo["tecnologias"])
        
        # 4. STACK TECNOLÓGICO
        try:
            stack_table = None

            # Buscar la tabla con la palabra clave "STACK TECNOLÓGICO" en la primera celda
            for table in doc.tables:
                if table.cell(0, 0).text.strip().lower() == "stack tecnológico":
                    stack_table = table
                break

            if not stack_table:
                print("⚠️ No se encontró la tabla de stack tecnológico.")
            else:
                print(f"✅ Tabla detectada con {len(stack_table.rows)} filas.")

                # Obtener lista de tecnologías
                tech_list = cv_data.get("stack_tecnologico", cv_data.get("stack", []))

                if not tech_list:
                    print("⚠️ No hay tecnologías para insertar.")
                else:
                    # Asegurarse de que haya al menos dos filas (título + contenido)
                    while len(stack_table.rows) < 2:
                        stack_table.add_row()

                    # Asegurarse de que la segunda fila tenga 2 celdas
                    while len(stack_table.rows[1].cells) < 2:
                        stack_table.rows[1]._tr.addnext(stack_table._tbl.tr_lst[-1])

                    # Limpiar las dos celdas de la segunda fila
                    for j in range(2):
                        stack_table.cell(1, j).text = ""

                    # Dividir tecnologías mitad y mitad
                    mid = (len(tech_list) + 1) // 2
                    col1 = tech_list[:mid]
                    col2 = tech_list[mid:]

                    # Insertar en la columna 0
                    cell1 = stack_table.cell(1, 0)
                    for tech in col1:
                        if tech and tech.lower() != "no especificado":
                            p = cell1.add_paragraph()
                            p.paragraph_format.tab_stops.add_tab_stop(Cm(0.64))
                            bullet= "\u25AA" + "\t"
                            p.add_run(bullet).bold = True
                            p.add_run(tech)
                    cell1.add_paragraph()

                    # Insertar en la columna 1
                    cell2 = stack_table.cell(1, 1)
                    for tech in col2:
                        if tech and tech.lower() != "no especificado":
                            p = cell2.add_paragraph()
                            p.paragraph_format.tab_stops.add_tab_stop(Cm(0.64))
                            bullet= "\u25AA" + "\t"
                            p.add_run(bullet).bold = True
                            p.add_run(tech)

                    print("✅ Sección stack tecnológico completada correctamente.")

        except Exception as e:
            print(f"❌ Error al procesar stack tecnológico: {str(e)}")

        
        # 5. FORMACIÓN ACADÉMICA
        # Encontrar los párrafos de formación académica
        formacion_paragraphs = []
        in_formacion = False
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == "FORMACIÓN ACADÉMICA":
                in_formacion = True
            elif para.text.strip() == "CURSOS Y CAPACITACIONES":
                in_formacion = False
            elif in_formacion and "Año, carrera de egreso, Institución" in para.text:
                formacion_paragraphs.append(i)
        
        # Reemplazar cada párrafo con la formación académica
        for i, para_idx in enumerate(formacion_paragraphs):
            if i < len(cv_data["formacion"]):
                replace_paragraph_text(doc.paragraphs[para_idx], cv_data["formacion"][i])
            else:
                replace_paragraph_text(doc.paragraphs[para_idx], "No especificado")
        
        # 6. CURSOS Y CAPACITACIONES
        # Encontrar los párrafos de cursos
        cursos_paragraphs = []
        in_cursos = False
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == "CURSOS Y CAPACITACIONES":
                in_cursos = True
            elif in_cursos and "Año (opcional), Curso, Institución" in para.text:
                cursos_paragraphs.append(i)
        
        # Reemplazar cada párrafo con los cursos
        for i, para_idx in enumerate(cursos_paragraphs):
            if i < len(cv_data["cursos"]):
                replace_paragraph_text(doc.paragraphs[para_idx], cv_data["cursos"][i])
            else:
                replace_paragraph_text(doc.paragraphs[para_idx], "No especificado")
        
        # Generar nombre único para el archivo
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        final_output_path = output_path.replace(".docx", f"_{timestamp}.docx")
        
        # Guardar documento
        doc.save(final_output_path)
        return final_output_path
    
    except Exception as e:
        print(f"Error al rellenar la plantilla: {str(e)}")
        traceback.print_exc()
        return None

# Función principal para procesar un CV (adaptada para Django)
def process_cv(input_file_path, template_path, api_key, model="o4-mini"):
    try:
        # Extraer texto del CV original
        _, file_extension = os.path.splitext(input_file_path)
        file_extension = file_extension.lower()
        
        if file_extension in ['.docx', '.doc']:
            cv_text = extract_text_from_word(input_file_path)
        elif file_extension == '.pdf':
            cv_text = extract_text_from_pdf(input_file_path)
        else:
            return {"success": False, "error": f"Formato de archivo no soportado: {file_extension}. Use PDF o DOCX."}
        
        if not cv_text:
            return {"success": False, "error": "No se pudo extraer texto del archivo proporcionado."}
        
        # Procesar con OpenAI
        response_text = process_with_openai(cv_text, api_key, model)
        
        if not response_text:
            return {"success": False, "error": "Error al procesar el CV con OpenAI"}
        
        # Extraer secciones
        cv_data = extract_sections(response_text)
        
        # Verificar datos mínimos
        if not cv_data["nombre"] or not cv_data["profesion"]:
            return {"success": False, "error": "No se pudo extraer información básica del CV (nombre o profesión)"}
        
        # Generar nombre de archivo de salida
        base_name = os.path.basename(input_file_path)
        name, _ = os.path.splitext(base_name)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(os.path.dirname(input_file_path), f"{name}_3IT_formato_{timestamp}.docx")
        
        # Aplicar datos a la plantilla
        result_path = fill_template_exact(cv_data, template_path, output_path)
        
        if result_path:
            return {"success": True, "output_path": result_path, "name": cv_data["nombre"]}
        else:
            return {"success": False, "error": "Error al generar el documento final."}
        
    except Exception as e:
        return {"success": False, "error": str(e)}
